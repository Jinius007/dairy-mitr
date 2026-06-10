"""One-off extractor: dumps the feed library xlsx and constraints docx to text."""
import json
import sys

OUT = r"d:\work\Innovations\AI Pashu Sahayak\c-Users-sinjini-Projects-dairy-sakha\scripts\extracted.txt"
XLSX = r"d:\work\Innovations\AI Pashu Sahayak\Knowledge Repository\Main Feed library .xlsx"
DOCX = r"d:\work\Innovations\AI Pashu Sahayak\Knowledge Repository\Other constraints used for RBP.docx"

lines = []

import openpyxl
wb = openpyxl.load_workbook(XLSX, data_only=True)
for ws in wb.worksheets:
    lines.append(f"===== SHEET: {ws.title} (dims {ws.dimensions}) =====")
    for row in ws.iter_rows(values_only=True):
        if all(c is None for c in row):
            continue
        vals = ["" if c is None else str(c) for c in row]
        lines.append(" | ".join(vals))

import docx
d = docx.Document(DOCX)
lines.append("===== DOCX: Other constraints used for RBP =====")
for p in d.paragraphs:
    if p.text.strip():
        lines.append(p.text)
for i, t in enumerate(d.tables):
    lines.append(f"--- DOCX TABLE {i} ---")
    for row in t.rows:
        lines.append(" | ".join(c.text.strip() for c in row.cells))

with open(OUT, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print("wrote", OUT, len(lines), "lines")
